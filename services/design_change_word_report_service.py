from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


class DesignChangeWordReportService:
    """설계변경·AI 품평 보고서 데이터를 다운로드용 DOCX로 변환합니다."""

    CHECK_LABELS = {
        "BOM_STRUCTURE": "BOM 구조",
        "LIFECYCLE": "Lifecycle",
        "APPROVAL": "자재 승인",
        "SUPPLIER": "공급사",
        "BOM_ATTRIBUTE": "BOM 속성",
        "COMPATIBILITY": "호환성",
    }
    STATUS_LABELS = {
        "PASS": "적합",
        "CONDITIONAL": "조건부 적합",
        "FAIL": "부적합",
        "NOT_CHECKED": "미검증",
    }

    @staticmethod
    def _value(value, default="-") -> str:
        if value is None or str(value).strip() in {"", "nan", "NaN"}:
            return default
        return str(value)

    @staticmethod
    def _set_cell_shading(cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)

    @staticmethod
    def _set_cell_width(cell, width_dxa: int) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.find(qn("w:tcW"))
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:w"), str(width_dxa))
        tc_w.set(qn("w:type"), "dxa")

    @staticmethod
    def _set_table_geometry(table, widths: list[int]) -> None:
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), str(sum(widths)))
        tbl_w.set(qn("w:type"), "dxa")
        tbl_ind = OxmlElement("w:tblInd")
        tbl_ind.set(qn("w:w"), "120")
        tbl_ind.set(qn("w:type"), "dxa")
        tbl_pr.append(tbl_ind)
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(width))
            grid.append(col)
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                DesignChangeWordReportService._set_cell_width(cell, width)

    @staticmethod
    def _set_font(run, size=10.5, bold=False, color="222222") -> None:
        run.font.name = "Malgun Gothic"
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Malgun Gothic")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)

    def _style_document(self, document: Document) -> None:
        section = document.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)

        normal = document.styles["Normal"]
        normal.font.name = "Malgun Gothic"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        normal.font.size = Pt(10.5)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.1

        for name, size, color, before, after in (
            ("Heading 1", 16, "2E74B5", 16, 8),
            ("Heading 2", 13, "2E74B5", 12, 6),
        ):
            style = document.styles[name]
            style.font.name = "Malgun Gothic"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)

        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._set_font(header.add_run("Display BOM AI Agent | Design Change Report"), 8.5, color="666666")
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_font(footer.add_run("AI 자동검증 결과는 사용자 최종 승인 전 양산 BOM에 반영되지 않습니다."), 8, color="777777")

    def _add_title(self, document: Document, report: dict) -> None:
        title = document.add_paragraph()
        title.paragraph_format.space_after = Pt(4)
        self._set_font(title.add_run("설계변경 · AI 품평 보고서"), 22, True, "1F3A5F")
        subtitle = document.add_paragraph()
        subtitle.paragraph_format.space_after = Pt(14)
        self._set_font(
            subtitle.add_run(
                f"변경 요청 {self._value(report.get('change_id'))}  |  "
                f"대상 모델 {self._value(report.get('product_id'))}"
            ),
            10.5,
            color="555555",
        )

    def _add_key_value_table(self, document: Document, rows: list[tuple[str, str]]) -> None:
        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = ""
            cells[1].text = ""
            self._set_cell_shading(cells[0], "E8EEF5")
            self._set_font(cells[0].paragraphs[0].add_run(label), 9.5, True, "1F3A5F")
            self._set_font(cells[1].paragraphs[0].add_run(self._value(value)), 9.5)
            for cell in cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        self._set_table_geometry(table, [2200, 7160])

    def _add_table(self, document: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            cell.text = ""
            self._set_cell_shading(cell, "E8EEF5")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_font(p.add_run(header), 8.5, True, "1F3A5F")
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = ""
                self._set_font(cells[index].paragraphs[0].add_run(self._value(value)), 8.5)
                cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        self._set_table_geometry(table, widths)

    def build(self, report: dict) -> bytes:
        if not report.get("success"):
            raise ValueError(report.get("message", "보고서 데이터 생성에 실패했습니다."))

        document = Document()
        self._style_document(document)
        self._add_title(document, report)

        change = report.get("change") or {}
        review = report.get("review") or {}
        document.add_heading("1. 설계변경 개요", level=1)
        self._add_key_value_table(document, [
            ("변경 요청 ID", report.get("change_id")),
            ("대상 모델", report.get("product_id")),
            ("변경 유형", change.get("change_type")),
            ("변경 사유", change.get("reason")),
            ("요청자", change.get("requested_by")),
            ("요청일 / 적용 예정일", f"{self._value(change.get('requested_date'))} / {self._value(change.get('effective_date'))}"),
            ("설계변경 분석 결과", self.STATUS_LABELS.get(str(change.get("analysis_result", "")).upper(), change.get("analysis_result"))),
        ])

        document.add_heading("2. 변경 전 · 후 BOM", level=1)
        changes = report.get("change_items") or []
        self._add_table(
            document,
            ["구분", "BOM Parent", "Location", "기존 자재", "신규 자재", "수량"],
            [[
                item.get("action"), item.get("bom_parent"), item.get("location"),
                item.get("old_bom_child"), item.get("new_bom_child"), item.get("quantity"),
            ] for item in changes],
            [900, 1700, 1350, 1900, 1900, 610],
        )

        document.add_heading("3. AI 품평 결과", level=1)
        review_result = str(review.get("review_result", "NOT_CHECKED")).upper()
        self._add_key_value_table(document, [
            ("Review BOM ID", review.get("review_id")),
            ("검증 Revision", report.get("report_revision")),
            ("AI 종합 판정", self.STATUS_LABELS.get(review_result, review_result)),
            ("검증 완료일", review.get("completed_date")),
            ("검증 주체", review.get("reviewed_by") or review.get("approved_by") or "BOM_AI_AGENT"),
        ])

        summary = report.get("review_check_summary") or {}
        summary_rows = []
        for check_type, values in summary.items():
            status = str(values.get("status", "NOT_CHECKED")).upper()
            summary_rows.append([
                self.CHECK_LABELS.get(check_type, check_type),
                self.STATUS_LABELS.get(status, status),
                values.get("count", 0),
                values.get("conditional_count", 0),
                values.get("fail_count", 0),
            ])
        self._add_table(
            document,
            ["검증 영역", "판정", "검증 수", "조건부", "실패"],
            summary_rows,
            [2600, 1900, 1620, 1620, 1620],
        )

        document.add_heading("4. 체크리스트 상세", level=1)
        checks = report.get("review_checks") or []
        check_rows = [[
            self.CHECK_LABELS.get(str(item.get("check_type", "")).upper(), item.get("check_type")),
            self.STATUS_LABELS.get(str(item.get("status", "")).upper(), item.get("status")),
            item.get("target_id"), item.get("message"),
        ] for item in checks]
        if check_rows:
            self._add_table(
                document,
                ["검증 영역", "판정", "대상", "검증 근거"],
                check_rows,
                [1700, 1200, 1900, 4560],
            )
        else:
            document.add_paragraph("저장된 체크리스트 상세 결과가 없습니다.")

        document.add_heading("5. 최종 권고", level=1)
        recommendation = (
            "AI 자동검증 항목이 모두 적합하므로 사용자 확인 후 양산 E-BOM 반영을 권고합니다."
            if review_result == "PASS"
            else "미검증·조건부·실패 항목을 확인하기 전에는 양산 E-BOM 반영을 권고하지 않습니다."
        )
        p = document.add_paragraph()
        self._set_font(p.add_run(recommendation), 11, True, "1F3A5F")
        note = document.add_paragraph()
        self._set_font(
            note.add_run("본 보고서는 PRE-APPLY 단계의 검토 자료이며, 사용자의 명시적 승인 전에는 Production BOM이 변경되지 않습니다."),
            9.5,
            color="666666",
        )

        output = BytesIO()
        document.save(output)
        return output.getvalue()
