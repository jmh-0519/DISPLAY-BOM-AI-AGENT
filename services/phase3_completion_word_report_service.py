from __future__ import annotations

import json
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


class Phase3CompletionWordReportService:
    """Build a business-grade final Phase3 design-change completion report.

    The active Phase3 process ends after final approval, Production E-BOM Apply and
    Word report generation.  STEP39 expands the report from a processing log into
    an evidence-oriented design-change completion document using persisted analysis,
    approval, preview, impact and apply data.
    """

    HEADER_FILL = "1F4E78"
    SUBHEADER_FILL = "D9EAF7"
    SUMMARY_FILL = "EAF3F8"
    SUCCESS_FILL = "E2F0D9"
    WARNING_FILL = "FFF2CC"
    BORDER_COLOR = "B7C9D6"

    @staticmethod
    def _is_blank(value: Any) -> bool:
        return value is None or value == "" or value == [] or value == {}

    @staticmethod
    def _text(value: Any) -> str:
        if Phase3CompletionWordReportService._is_blank(value):
            return "-"
        if isinstance(value, bool):
            return "Y" if value else "N"
        return str(value)

    @classmethod
    def _number(cls, value: Any, digits: int = 2) -> str:
        if cls._is_blank(value):
            return "-"
        try:
            number = float(value)
        except (TypeError, ValueError):
            return cls._text(value)
        if number.is_integer():
            return f"{int(number):,}"
        return f"{number:,.{digits}f}".rstrip("0").rstrip(".")

    @classmethod
    def _money(cls, value: Any, currency: Any = None) -> str:
        if cls._is_blank(value):
            return "-"
        suffix = f" {currency}" if currency not in {None, ""} else ""
        return f"{cls._number(value)}{suffix}"

    @staticmethod
    def _font(run, size: float = 9.5, bold: bool = False, color: str | None = None):
        run.font.name = "Malgun Gothic"
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)

    @staticmethod
    def _shade(cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)

    @staticmethod
    def _set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_mar = tc_pr.first_child_found_in("w:tcMar")
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
            node = tc_mar.find(qn(f"w:{margin}"))
            if node is None:
                node = OxmlElement(f"w:{margin}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    def _table(
        self,
        document: Document,
        headers: list[str],
        rows: list[list[Any]],
        *,
        widths: list[float] | None = None,
        header_fill: str | None = None,
    ):
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = True
        header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        header_tr_pr.append(tbl_header)
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            self._set_cell_margins(cell)
            self._shade(cell, header_fill or self.HEADER_FILL)
            run = cell.paragraphs[0].add_run(header)
            self._font(run, bold=True, color="FFFFFF")
            if widths and idx < len(widths):
                cell.width = Cm(widths[idx])
        for row in rows:
            table_row = table.add_row()
            tr_pr = table_row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)
            cells = table_row.cells
            for idx, value in enumerate(row):
                cells[idx].text = ""
                cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                self._set_cell_margins(cells[idx])
                self._font(cells[idx].paragraphs[0].add_run(self._text(value)), 9)
                if widths and idx < len(widths):
                    cells[idx].width = Cm(widths[idx])
        document.add_paragraph().paragraph_format.space_after = Pt(1)
        return table

    def _key_value_table(self, document: Document, rows: list[list[Any]]):
        table = self._table(document, ["항목", "내용"], rows, widths=[4.0, 12.5])
        for row in table.rows[1:]:
            self._shade(row.cells[0], self.SUBHEADER_FILL)
            for run in row.cells[0].paragraphs[0].runs:
                run.bold = True
        return table

    def _section(self, document: Document, title: str, level: int = 1):
        paragraph = document.add_heading(title, level=level)
        paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
        paragraph.paragraph_format.space_after = Pt(5)
        for run in paragraph.runs:
            self._font(run, 13 if level == 1 else 11, True, "1F1F1F")
        return paragraph

    def _paragraph(self, document: Document, text: str, *, bold_prefix: str | None = None):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(5)
        if bold_prefix and text.startswith(bold_prefix):
            first, rest = text[: len(bold_prefix)], text[len(bold_prefix):]
            self._font(paragraph.add_run(first), 9.5, True)
            self._font(paragraph.add_run(rest), 9.5)
        else:
            self._font(paragraph.add_run(text), 9.5)
        return paragraph

    def _summary_box(self, document: Document, text: str, *, fill: str | None = None):
        table = document.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = table.cell(0, 0)
        self._shade(cell, fill or self.SUMMARY_FILL)
        self._set_cell_margins(cell, top=130, start=150, bottom=130, end=150)
        cell.text = ""
        self._font(cell.paragraphs[0].add_run(text), 10, True)
        document.add_paragraph().paragraph_format.space_after = Pt(1)

    @staticmethod
    def _reason_evidence_text(evidence: Any) -> str:
        if not evidence:
            return "-"
        if not isinstance(evidence, dict):
            return str(evidence)
        labels = {
            "input_source": "입력 근거",
            "raw_reason_text": "사용자 표현",
            "db_lifecycle_status": "DB Lifecycle",
            "lifecycle_status": "Lifecycle",
            "reason": "판단 사유",
            "message": "메시지",
            "description": "설명",
            "alias_matches": "인식 표현",
            "all_detected_reason_codes": "감지 사유",
            "primary_reason_code": "Primary Reason",
        }
        preferred = []
        for key, label in labels.items():
            value = evidence.get(key)
            if Phase3CompletionWordReportService._is_blank(value):
                continue
            if isinstance(value, list):
                value = ", ".join(str(item) for item in value)
            preferred.append(f"{label}: {value}")
        if preferred:
            return "; ".join(preferred)
        compact = {
            key: value for key, value in evidence.items()
            if not Phase3CompletionWordReportService._is_blank(value)
        }
        if not compact:
            return "-"
        return json.dumps(compact, ensure_ascii=False, separators=(", ", ": "))

    @staticmethod
    def _candidate_counts(analysis_summary: dict) -> str:
        counts = analysis_summary.get("status_counts") or {}
        return (
            f"총 {analysis_summary.get('candidate_count') or 0}건 "
            f"(PASS {counts.get('PASS', 0)} / CONDITIONAL {counts.get('CONDITIONAL', 0)} / "
            f"FAIL {counts.get('FAIL', 0)})"
        )

    def _executive_summary(self, report: dict) -> str:
        request = report.get("request") or {}
        actions = report.get("actions") or []
        analysis_summary = report.get("analysis_summary") or {}
        impact = report.get("impact_review") or {}
        return (
            f"{request.get('version_code') or '-'} / {request.get('plant_code') or '-'}의 설계변경 "
            f"{len(actions)}건이 후보 분석과 영향 검토, 최종 승인을 거쳐 Production E-BOM에 적용되었습니다. "
            f"후보 검토 결과는 {self._candidate_counts(analysis_summary)}이며, "
            f"영향 모델은 {impact.get('impacted_model_count') or 0}개입니다. "
            f"최종 Apply 결과는 {(report.get('apply_result') or {}).get('result') or '-'}입니다."
        )

    def build(self, report: dict) -> bytes:
        if not report.get("success"):
            raise ValueError(report.get("message") or "설계변경 완료 보고서 데이터를 생성할 수 없습니다.")

        document = Document()
        section = document.sections[0]
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.7)
        section.right_margin = Cm(1.7)

        normal = document.styles["Normal"]
        normal.font.name = "Malgun Gothic"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        normal.font.size = Pt(9.5)

        request = report.get("request") or {}
        actions = report.get("actions") or []
        analysis_summary = report.get("analysis_summary") or {}
        selected_details = report.get("selected_candidate_details") or []
        impact_review = report.get("impact_review") or {}
        approvals = report.get("approvals") or []
        preview = report.get("preview") or {}
        apply_result = report.get("apply_result") or {}

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(2)
        self._font(title.add_run("Display BOM AI Agent"), 18, True, self.HEADER_FILL)
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(10)
        self._font(subtitle.add_run("설계변경 완료 보고서"), 15, True)

        meta = document.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._font(
            meta.add_run(
                f"Request ID: {self._text(report.get('request_id'))}  |  "
                f"PLANT: {self._text(request.get('plant_code'))}  |  "
                f"제품: {self._text(request.get('version_code'))}"
            ),
            9,
            False,
            "666666",
        )

        self._section(document, "1. 완료 요약")
        self._summary_box(document, self._executive_summary(report), fill=self.SUCCESS_FILL)
        selected_statuses = [detail.get("final_status") for detail in selected_details]
        self._table(document, ["구분", "결과"], [
            ["업무 상태", "COMPLETED"],
            ["Production Apply", apply_result.get("result")],
            ["변경 Action", f"{len(actions)}건"],
            ["후보 검토", self._candidate_counts(analysis_summary)],
            ["최종 선택 후보", f"{len(selected_details)}건 / " + ", ".join(selected_statuses or ["-"])],
            ["영향 모델", f"{impact_review.get('impacted_model_count') or 0}개"],
        ], widths=[4.2, 12.3])

        self._section(document, "2. 설계변경 개요")
        self._key_value_table(document, [
            ["Request ID", report.get("request_id")],
            ["PLANT", request.get("plant_code")],
            ["제품", request.get("version_code")],
            ["요청 원문", request.get("original_request")],
            ["변경 사유", ", ".join(request.get("reasons") or [])],
            ["기준일", request.get("as_of_date")],
            ["적용일", request.get("effective_date")],
            ["요청자", request.get("requested_by")],
            ["최종 Workflow 상태", request.get("workflow_status")],
            ["Apply 상태", request.get("apply_status")],
        ])

        self._section(document, "3. 변경 전 / 후 확정 내용")
        before_after_rows = []
        for action in actions:
            before_label = " / ".join(filter(None, [
                action.get("old_item_code"), action.get("old_item_name"), action.get("old_item_description")
            ])) or "-"
            after_label = " / ".join(filter(None, [
                action.get("new_item_code"), action.get("new_item_name"), action.get("new_item_description")
            ])) or "-"
            parent_label = " / ".join(filter(None, [action.get("parent_item_code"), action.get("parent_item_name")]))
            old_quantity = action.get("old_quantity")
            new_quantity = action.get("new_quantity")
            if new_quantity is None and action.get("action_type") == "REPLACE":
                new_quantity = old_quantity
            quantity = f"{self._number(old_quantity)} → {self._number(new_quantity)}"
            before_after_rows.append([
                action.get("action_seq"), action.get("action_type"), parent_label,
                f"{before_label}  →  {after_label}",
                f"{action.get('location_code') or '-'} / 수량 {quantity}",
                action.get("evaluation_status"),
            ])
        self._table(
            document,
            ["순번", "Action", "Parent", "변경 전 → 변경 후", "Location / 수량", "평가"],
            before_after_rows,
        )

        self._section(document, "4. 변경 사유 및 Evidence")
        has_reason = False
        for action in actions:
            for reason in action.get("reasons") or []:
                has_reason = True
                label = "Primary" if reason.get("is_primary") == "Y" else "Secondary"
                self._section(
                    document,
                    f"Action {action.get('action_seq') or '-'} · {label} Reason: {reason.get('reason_code') or '-'}",
                    level=2,
                )
                self._key_value_table(document, [
                    ["사용자 표현", reason.get("raw_reason_text")],
                    ["해석 상태", reason.get("resolution_status")],
                    ["해석 근거", reason.get("resolution_source")],
                    ["Confidence", self._number(reason.get("confidence"))],
                    ["Evidence", self._reason_evidence_text(reason.get("evidence"))],
                ])
        if not has_reason:
            self._paragraph(document, "저장된 변경사유 Evidence가 없습니다.")

        self._section(document, "5. 후보 분석 및 최종 선정 근거")
        self._summary_box(document, analysis_summary.get("summary") or self._candidate_counts(analysis_summary))
        candidate_rows = []
        selected_ids = {action.get("selected_candidate_id") for action in actions}
        action_seq_by_id = {action.get("action_id"): action.get("action_seq") for action in actions}
        for candidate in report.get("candidate_evaluations") or []:
            candidate_rows.append([
                action_seq_by_id.get(candidate.get("action_id")),
                candidate.get("candidate_item_code"),
                candidate.get("candidate_item_name"),
                candidate.get("final_status"),
                f"{self._number(candidate.get('total_score'))} / {self._text(candidate.get('grade'))}",
                candidate.get("rank_no"),
                "Y" if candidate.get("candidate_id") in selected_ids else "",
            ])
        if candidate_rows:
            self._table(
                document,
                ["Action", "후보 코드", "후보명", "종합", "점수/등급", "순위", "최종선정"],
                candidate_rows,
            )

        for detail in selected_details:
            candidate_item = detail.get("candidate_item") or {}
            source_item = detail.get("source_item") or {}
            self._section(
                document,
                f"Action {detail.get('action_seq') or '-'} 최종 후보: {candidate_item.get('item_code') or '-'}",
                level=2,
            )
            self._key_value_table(document, [
                ["기존 품목", " / ".join(filter(None, [source_item.get("item_code"), source_item.get("item_name"), source_item.get("description")]))],
                ["선정 후보", " / ".join(filter(None, [candidate_item.get("item_code"), candidate_item.get("item_name"), candidate_item.get("description")]))],
                ["종합 적합성", detail.get("final_status")],
                ["점수 / 등급 / 순위", f"{self._number(detail.get('total_score'))} / {self._text(detail.get('grade'))} / {self._text(detail.get('rank'))}"],
                ["Primary Reason", (detail.get("reason_context") or {}).get("primary_reason")],
                ["Secondary Reasons", ", ".join((detail.get("reason_context") or {}).get("secondary_reasons") or [])],
                ["누락 데이터", ", ".join(detail.get("missing_data") or [])],
            ])

        self._section(document, "6. 기술 적합성 검증")
        technical_rows = []
        technical_explanations = []
        for detail in selected_details:
            candidate_code = (detail.get("candidate_item") or {}).get("item_code")
            tech = detail.get("technical_evaluation") or {}
            checks = tech.get("checks") or []
            if not checks:
                technical_rows.append([candidate_code, "-", "-", "-", "-", tech.get("status")])
            for check in checks:
                mode_rule = " / ".join(filter(None, [check.get("evaluation_mode"), check.get("rule_id")])) or "-"
                standard = self._text(check.get("candidate_value"))
                if not self._is_blank(check.get("expected_value")):
                    standard += f" / 기준 {self._text(check.get('expected_value'))} ({self._text(check.get('operator'))})"
                technical_rows.append([
                    candidate_code, mode_rule, check.get("attribute"),
                    check.get("source_value"), standard, check.get("status"),
                ])
            explanation = tech.get("explanation")
            if explanation:
                if isinstance(explanation, list):
                    explanation = " ".join(str(item) for item in explanation)
                technical_explanations.append(f"{candidate_code}: {explanation}")
        if technical_rows:
            self._table(
                document,
                ["후보", "평가방식 / Rule", "항목", "변경 전", "후보 / 기준", "결과"],
                technical_rows,
            )
            for text in technical_explanations:
                self._paragraph(document, f"판단 요약 · {text}")
        else:
            self._paragraph(document, "저장된 기술 검증 상세 Evidence가 없습니다.")

        self._section(document, "7. 공급사 및 원가 평가")
        if selected_details:
            for detail in selected_details:
                supplier = detail.get("supplier_evaluation") or {}
                candidate_code = (detail.get("candidate_item") or {}).get("item_code")
                self._section(document, f"후보 {candidate_code or '-'}", level=2)
                self._key_value_table(document, [
                    ["공급 평가", supplier.get("status")],
                    ["선정 공급사", " / ".join(filter(None, [supplier.get("supplier_code"), supplier.get("supplier_name")]))],
                    ["후보 단가", self._money(supplier.get("unit_price"), supplier.get("currency_code"))],
                    ["납기", f"{self._number(supplier.get('lead_time_days'))}일" if supplier.get("lead_time_days") is not None else "-"],
                    ["품질 등급", supplier.get("quality_grade")],
                    ["공급 안정성", self._number(supplier.get("stability_score"))],
                    ["공급 상태", supplier.get("supply_status")],
                    ["선정 근거", supplier.get("decision_reason")],
                ])
            self._paragraph(
                document,
                "※ 기존 품목과 후보 품목의 양쪽 단가 Evidence가 저장된 경우에만 원가 절감액을 확정할 수 있습니다. 후보 단가만 존재하는 경우 절감액은 임의 계산하지 않습니다.",
            )
        else:
            self._paragraph(document, "저장된 공급사/원가 평가 정보가 없습니다.")

        self._section(document, "8. BOM 수량 및 재고 검증")
        if selected_details:
            for detail in selected_details:
                inventory = detail.get("inventory_evaluation") or {}
                demand_context = inventory.get("demand_context") or {}
                candidate_code = (detail.get("candidate_item") or {}).get("item_code")
                self._section(document, f"후보 {candidate_code or '-'}", level=2)
                self._key_value_table(document, [
                    ["재고 평가", inventory.get("status")],
                    ["BOM 수량", self._number(demand_context.get("bom_quantity") if demand_context.get("bom_quantity") is not None else inventory.get("demand_quantity") or demand_context.get("quantity"))],
                    ["가용재고", self._number(inventory.get("available_quantity"))],
                    ["부족수량", self._number(inventory.get("shortage_quantity"))],
                    ["적용일", inventory.get("effective_date") or request.get("effective_date")],
                ])
                explanations = inventory.get("explanation") or []
                if isinstance(explanations, str):
                    explanations = [explanations]
                for text in explanations:
                    self._paragraph(document, f"판단 근거 · {text}")
        else:
            self._paragraph(document, "저장된 BOM 수량/재고 검증 정보가 없습니다.")

        self._section(document, "9. BOM 영향 분석")
        if impact_review:
            self._key_value_table(document, [
                ["공용 BOM 영향 승인 필요", impact_review.get("requires_impact_approval")],
                ["영향 모델 수", impact_review.get("impacted_model_count")],
                ["PLANT", impact_review.get("plant_code")],
            ])
            for idx, action in enumerate(impact_review.get("actions") or [], 1):
                model_codes = ", ".join(
                    row.get("model_code") for row in action.get("impacted_models") or [] if row.get("model_code")
                )
                self._section(document, f"영향 Action {idx}", level=2)
                self._key_value_table(document, [
                    ["Parent", action.get("parent_item_code")],
                    ["공용/단독", action.get("parent_usage_type")],
                    ["공용 BOM 변경", action.get("shared_bom_change")],
                    ["영향 모델", model_codes or "-"],
                    ["변경 Spec 수", action.get("changed_spec_count")],
                ])
                changed_specs = action.get("changed_specs") or []
                if changed_specs:
                    self._table(
                        document,
                        ["Spec 항목", "변경 전", "변경 후", "변화"],
                        [[row.get("attribute"), row.get("before"), row.get("after"), row.get("change_status")] for row in changed_specs],
                    )
        else:
            snapshot_impacts = (preview.get("snapshot") or {}).get("impacts") or []
            if snapshot_impacts:
                self._table(
                    document,
                    ["Action", "영향 품목", "영향유형", "영향 경로"],
                    [[row.get("action_id"), row.get("impacted_item_code"), row.get("impact_type"), row.get("impact_path")] for row in snapshot_impacts],
                )
            else:
                self._paragraph(document, "저장된 BOM 영향 분석 결과가 없습니다.")

        self._section(document, "10. 승인 이력 및 최종 Preview")
        approval_rows = []
        for approval in approvals:
            approval_rows.append([
                approval.get("approval_stage"), approval.get("decision"), approval.get("approved_by"),
                approval.get("decision_reason"), approval.get("approved_at"),
            ])
        if approval_rows:
            self._table(document, ["승인 단계", "결정", "승인자", "사유", "승인시각"], approval_rows)
        self._key_value_table(document, [
            ["Preview ID", preview.get("preview_id")],
            ["Revision", preview.get("preview_revision")],
            ["검증 상태", preview.get("validation_status")],
            ["생성자", preview.get("created_by")],
            ["생성시각", preview.get("created_at")],
        ])

        snapshot_actions = (preview.get("snapshot") or {}).get("actions") or []
        if snapshot_actions:
            seq_by_id = {action.get("action_id"): action.get("action_seq") for action in actions}
            self._table(
                document,
                ["Action", "변경 전", "변경 후", "원본 BOM 수량", "변경 수량", "평가"],
                [[
                    seq_by_id.get(row.get("action_id")), row.get("old_item_code"), row.get("new_item_code"),
                    self._number(row.get("source_bom_quantity")), self._number(row.get("new_quantity")),
                    row.get("evaluation_status"),
                ] for row in snapshot_actions],
            )

        self._section(document, "11. Production E-BOM 적용 결과")
        action_results = apply_result.get("action_results") or []
        self._key_value_table(document, [
            ["Apply ID", apply_result.get("apply_id")],
            ["결과", apply_result.get("result")],
            ["적용 Action 수", len(action_results) or len(actions)],
            ["적용자", apply_result.get("applied_by")],
            ["적용시각", apply_result.get("created_at")],
            ["Preview ID", apply_result.get("preview_id")],
            ["최종 승인 ID", apply_result.get("final_approval_id")],
        ])

        self._section(document, "12. 최종 결론")
        conclusion = (
            f"Request {report.get('request_id')}의 설계변경 {len(actions)}건은 최종 승인 후 "
            f"Production E-BOM에 {apply_result.get('result') or '-'} 상태로 반영되었습니다. "
            "본 보고서는 설계변경 요청, 후보 평가 Evidence, BOM 영향, 승인 및 Apply 결과를 동일 Request 기준으로 기록합니다."
        )
        self._summary_box(document, conclusion, fill=self.SUCCESS_FILL)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._font(
            footer.add_run(
                f"Display BOM AI Agent | {self._text(report.get('request_id'))} | Completion Report"
            ),
            8,
            False,
            "777777",
        )

        stream = BytesIO()
        document.save(stream)
        return stream.getvalue()
