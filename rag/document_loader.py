from __future__ import annotations

import re
import tomllib
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any, Iterable

from .knowledge_models import KnowledgeDocument, KnowledgeDocumentMetadata, KnowledgeSection


_ALLOWED_DOCUMENT_TYPES = {
    "CHANGE_RULE",
    "CHANGE_REASON",
    "DESIGN_GUIDE",
    "MATERIAL_SPEC",
    "PROCESS_GUIDE",
    "CHANGE_POLICY",
    "SUPPLIER_TECHNICAL",
    "FAQ",
}
_ALLOWED_STATUS = {"ACTIVE", "INACTIVE", "DRAFT"}
_SUPPORTED_EXTENSIONS = {".md", ".txt", ".docx", ".pdf"}


class KnowledgeDocumentError(ValueError):
    """A knowledge document cannot be loaded under the managed RAG contract."""


def _split_toml_front_matter(text: str, path: Path) -> tuple[dict[str, Any], str]:
    lines = text.splitlines()
    if not lines or lines[0].strip() != "+++":
        raise KnowledgeDocumentError(f"{path}: TOML front matter must start with +++")
    try:
        end_index = next(
            index for index, line in enumerate(lines[1:], 1) if line.strip() == "+++"
        )
    except StopIteration as exc:
        raise KnowledgeDocumentError(f"{path}: TOML front matter closing +++ not found") from exc
    try:
        metadata = tomllib.loads("\n".join(lines[1:end_index]).strip())
    except tomllib.TOMLDecodeError as exc:
        raise KnowledgeDocumentError(f"{path}: invalid TOML front matter: {exc}") from exc
    return metadata, "\n".join(lines[end_index + 1 :]).strip()


def _validate_date(value: str, field_name: str) -> None:
    try:
        date.fromisoformat(value)
    except ValueError as exc:
        raise KnowledgeDocumentError(f"{field_name} must use YYYY-MM-DD: {value}") from exc


def _tuple_of_strings(value: Any, *, field_name: str) -> tuple[str, ...]:
    if value is None:
        return ()
    if not isinstance(value, list):
        raise KnowledgeDocumentError(f"{field_name} must be a list")
    return tuple(str(item).strip() for item in value if str(item).strip())


def _normalize_metadata(raw: dict[str, Any], *, source_path: Path) -> KnowledgeDocumentMetadata:
    document_id = str(raw.get("document_id") or "").strip()
    document_title = str(raw.get("document_title") or "").strip()
    document_type = str(raw.get("document_type") or "").strip().upper()
    version = str(raw.get("version") or "").strip()
    effective_date = str(raw.get("effective_date") or "").strip()
    status = str(raw.get("status") or "ACTIVE").strip().upper()
    language = str(raw.get("language") or "KO").strip().upper()

    required = {
        "document_id": document_id,
        "document_title": document_title,
        "document_type": document_type,
        "version": version,
        "effective_date": effective_date,
    }
    missing = [name for name, value in required.items() if not value]
    if missing:
        raise KnowledgeDocumentError(f"{source_path}: missing metadata fields {missing}")
    if document_type not in _ALLOWED_DOCUMENT_TYPES:
        raise KnowledgeDocumentError(
            f"{source_path}: unsupported document_type {document_type}; "
            f"allowed={sorted(_ALLOWED_DOCUMENT_TYPES)}"
        )
    if status not in _ALLOWED_STATUS:
        raise KnowledgeDocumentError(f"{source_path}: invalid status {status}")
    _validate_date(effective_date, f"{source_path}: effective_date")

    attributes = raw.get("attributes") or {}
    if not isinstance(attributes, dict):
        raise KnowledgeDocumentError(f"{source_path}: attributes must be a TOML table")

    return KnowledgeDocumentMetadata(
        document_id=document_id,
        document_title=document_title,
        document_type=document_type,
        version=version,
        effective_date=effective_date,
        status=status,
        language=language,
        source_path=source_path,
        product_families=_tuple_of_strings(
            raw.get("product_families"), field_name=f"{source_path}: product_families"
        ),
        material_types=_tuple_of_strings(
            raw.get("material_types"), field_name=f"{source_path}: material_types"
        ),
        tags=_tuple_of_strings(raw.get("tags"), field_name=f"{source_path}: tags"),
        attributes=dict(attributes),
    )


def _markdown_sections(text: str, *, fallback_title: str) -> tuple[KnowledgeSection, ...]:
    heading_stack: list[str] = []
    current_title = fallback_title
    current_path: tuple[str, ...] = (fallback_title,)
    buffer: list[str] = []
    sections: list[KnowledgeSection] = []

    def flush() -> None:
        content = "\n".join(buffer).strip()
        if content:
            sections.append(
                KnowledgeSection(
                    title=current_title,
                    path=current_path,
                    content=content,
                    order=len(sections) + 1,
                )
            )
        buffer.clear()

    for line in text.splitlines():
        match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if not match:
            buffer.append(line)
            continue
        flush()
        level = len(match.group(1))
        title = match.group(2).strip()
        heading_stack[:] = heading_stack[: level - 1]
        while len(heading_stack) < level - 1:
            heading_stack.append(fallback_title)
        heading_stack.append(title)
        current_title = title
        current_path = tuple(heading_stack)

    flush()
    if not sections:
        sections.append(
            KnowledgeSection(
                title=fallback_title,
                path=(fallback_title,),
                content=text.strip(),
                order=1,
            )
        )
    return tuple(sections)


def _plain_text_sections(text: str, *, title: str) -> tuple[KnowledgeSection, ...]:
    return (
        KnowledgeSection(title=title, path=(title,), content=text.strip(), order=1),
    )


def _sidecar_path(path: Path) -> Path:
    return path.with_name(f"{path.stem}.meta.toml")


@dataclass(frozen=True)
class KnowledgeDocumentLoader:
    """Load managed knowledge documents without performing embedding or retrieval."""

    def load(self, path: str | Path) -> KnowledgeDocument:
        source_path = Path(path)
        suffix = source_path.suffix.lower()
        if suffix not in _SUPPORTED_EXTENSIONS:
            raise KnowledgeDocumentError(f"unsupported knowledge document type: {source_path}")
        if not source_path.exists():
            raise KnowledgeDocumentError(f"knowledge document not found: {source_path}")

        if suffix == ".md":
            metadata_raw, body = _split_toml_front_matter(
                source_path.read_text(encoding="utf-8"), source_path
            )
            metadata = _normalize_metadata(metadata_raw, source_path=source_path)
            sections = _markdown_sections(body, fallback_title=metadata.document_title)
        else:
            metadata_path = _sidecar_path(source_path)
            if not metadata_path.exists():
                raise KnowledgeDocumentError(
                    f"{source_path}: metadata sidecar is required: {metadata_path.name}"
                )
            try:
                metadata_raw = tomllib.loads(metadata_path.read_text(encoding="utf-8"))
            except tomllib.TOMLDecodeError as exc:
                raise KnowledgeDocumentError(
                    f"{metadata_path}: invalid TOML metadata: {exc}"
                ) from exc
            metadata = _normalize_metadata(metadata_raw, source_path=source_path)
            if suffix == ".txt":
                sections = _plain_text_sections(
                    source_path.read_text(encoding="utf-8"), title=metadata.document_title
                )
            elif suffix == ".docx":
                sections = self._load_docx_sections(source_path, metadata.document_title)
            else:
                sections = self._load_pdf_sections(source_path, metadata.document_title)

        if not any(section.content.strip() for section in sections):
            raise KnowledgeDocumentError(f"{source_path}: document has no extractable text")
        return KnowledgeDocument(metadata=metadata, sections=sections)

    def load_directory(self, path: str | Path) -> tuple[KnowledgeDocument, ...]:
        root = Path(path)
        if not root.exists():
            return ()
        documents: list[KnowledgeDocument] = []
        for source_path in sorted(root.rglob("*")):
            if not source_path.is_file():
                continue
            if source_path.name.lower() == "readme.md":
                continue
            if source_path.name.endswith(".meta.toml"):
                continue
            if source_path.suffix.lower() not in _SUPPORTED_EXTENSIONS:
                continue
            documents.append(self.load(source_path))
        return tuple(documents)

    @staticmethod
    def _load_docx_sections(path: Path, fallback_title: str) -> tuple[KnowledgeSection, ...]:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise KnowledgeDocumentError(
                "python-docx is required to load .docx knowledge documents"
            ) from exc

        document = DocxDocument(path)
        sections: list[KnowledgeSection] = []
        heading_stack: list[str] = []
        current_title = fallback_title
        current_path: tuple[str, ...] = (fallback_title,)
        buffer: list[str] = []

        def flush() -> None:
            content = "\n".join(buffer).strip()
            if content:
                sections.append(
                    KnowledgeSection(
                        title=current_title,
                        path=current_path,
                        content=content,
                        order=len(sections) + 1,
                    )
                )
            buffer.clear()

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = str(getattr(paragraph.style, "name", "") or "")
            heading_match = re.match(r"Heading\s+(\d+)", style_name, flags=re.IGNORECASE)
            if not heading_match:
                buffer.append(text)
                continue
            flush()
            level = max(1, int(heading_match.group(1)))
            heading_stack[:] = heading_stack[: level - 1]
            while len(heading_stack) < level - 1:
                heading_stack.append(fallback_title)
            heading_stack.append(text)
            current_title = text
            current_path = tuple(heading_stack)

        flush()
        if not sections:
            combined = "\n".join(p.text.strip() for p in document.paragraphs if p.text.strip())
            return _plain_text_sections(combined, title=fallback_title)
        return tuple(sections)

    @staticmethod
    def _load_pdf_sections(path: Path, fallback_title: str) -> tuple[KnowledgeSection, ...]:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise KnowledgeDocumentError(
                "pypdf is required to load text-based .pdf knowledge documents"
            ) from exc

        reader = PdfReader(path)
        sections: list[KnowledgeSection] = []
        for index, page in enumerate(reader.pages, 1):
            content = str(page.extract_text() or "").strip()
            if not content:
                continue
            page_title = f"Page {index}"
            sections.append(
                KnowledgeSection(
                    title=page_title,
                    path=(fallback_title, page_title),
                    content=content,
                    order=len(sections) + 1,
                    page_number=index,
                )
            )
        return tuple(sections)


__all__ = [
    "KnowledgeDocumentError",
    "KnowledgeDocumentLoader",
]
